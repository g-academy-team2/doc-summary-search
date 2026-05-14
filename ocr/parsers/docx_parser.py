import logging
from docx import Document
from .vision_helper import get_text_from_image_bytes

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        
        for para in doc.paragraphs:
            if para.text.strip(): 
                full_text.append(para.text)
        
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text)) 
        
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                try:
                    logger.info(f"📸 {file_path} 내 {image_count}번째 이미지 분석 중...")
                    image_bytes = rel.target_part.blob
                    vision_text = get_text_from_image_bytes(image_bytes)
                    if vision_text:
                        full_text.append(f"\n[문서 내 이미지 OCR]:\n{vision_text}")
                except Exception as img_err:
                    logger.warning(f"⚠️ 이미지 추출 중 건너뜀: {img_err}")
                
        return "\n\n".join(full_text)

    except Exception as e:
        logger.error(f"❌ DOCX 파싱 에러 ({file_path}): {e}")
        return f"문서 읽기 실패: {str(e)}"