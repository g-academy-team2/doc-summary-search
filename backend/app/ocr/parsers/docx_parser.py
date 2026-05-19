import os
import logging
import uuid # 💡 임시 파일 이름을 안전하게 만들기 위해 추가
from docx import Document

# 💡 1. 새로운 비전 헬퍼 함수명과 절대 경로로 수정했습니다.
from app.ocr.parsers.vision_helper import extract_text_from_image

logger = logging.getLogger(__name__)

def extract_text_from_docx(file_path):
    try:
        doc = Document(file_path)
        full_text = []
        
        # 단락(Paragraph) 텍스트 추출
        for para in doc.paragraphs:
            if para.text.strip(): 
                full_text.append(para.text)
        
        # 표(Table) 텍스트 추출
        for table in doc.tables:
            for row in table.rows:
                row_text = [cell.text.strip() for cell in row.cells if cell.text.strip()]
                if row_text:
                    full_text.append(" | ".join(row_text)) 
        
        # 문서 내 이미지 추출 및 OCR 처리
        image_count = 0
        for rel in doc.part.rels.values():
            if "image" in rel.target_ref:
                image_count += 1
                try:
                    logger.info(f"📸 {file_path} 내 {image_count}번째 이미지 분석 중...")
                    image_bytes = rel.target_part.blob
                    
                    # 💡 2. 바이트 데이터를 임시 이미지 파일로 저장합니다.
                    temp_img_path = f"temp_docx_img_{uuid.uuid4().hex}.png"
                    with open(temp_img_path, "wb") as f:
                        f.write(image_bytes)
                    
                    # 💡 3. 실제 파일 경로를 넘겨서 비전 헬퍼에게 읽게 합니다.
                    vision_text = extract_text_from_image(temp_img_path)
                    
                    # 💡 4. 다 쓴 임시 파일은 하드디스크 용량을 차지하지 않게 즉시 삭제!
                    if os.path.exists(temp_img_path):
                        os.remove(temp_img_path)
                        
                    if vision_text:
                        full_text.append(f"\n[문서 내 이미지 OCR]:\n{vision_text}")
                        
                except Exception as img_err:
                    logger.warning(f"⚠️ 이미지 추출 중 건너뜀: {img_err}")
                
        return "\n\n".join(full_text)

    except Exception as e:
        logger.error(f"❌ DOCX 파싱 에러 ({file_path}): {e}")
        return f"문서 읽기 실패: {str(e)}"